"""
Script para verificar los placeholders en el template Word
"""
from docxtpl import DocxTemplate
import re
import os

def verificar_placeholders():
    template_path = os.path.join(os.path.dirname(__file__), 'templates', 'resolucion_template.docx')
    
    if not os.path.exists(template_path):
        print(f"[ERROR] No se encontro el template en: {template_path}")
        return
    
    print(f"[OK] Template encontrado: {template_path}\n")
    
    try:
        # Cargar el template
        doc = DocxTemplate(template_path)
        
        # Extraer texto del documento usando python-docx directamente
        from docx import Document
        docx_file = Document(template_path)
        
        texto_completo = ""
        
        # Extraer de párrafos
        for paragraph in docx_file.paragraphs:
            texto_completo += paragraph.text + "\n"
            # También buscar en runs individuales (pueden tener formato diferente)
            for run in paragraph.runs:
                if run.text:
                    texto_completo += run.text + "\n"
        
        # También buscar en tablas
        for table in docx_file.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + "\n"
                    # También en los párrafos de las celdas
                    for paragraph in cell.paragraphs:
                        texto_completo += paragraph.text + "\n"
        
        # Buscar placeholders con {{ }}
        placeholders_doble = re.findall(r'\{\{([^}]+)\}\}', texto_completo)
        
        # Buscar placeholders con {$ }
        placeholders_dolar = re.findall(r'\{\$([^}]+)\}', texto_completo)
        
        # Buscar placeholders con {% %}
        placeholders_loop = re.findall(r'\{%\s*for\s+(\w+)\s+in\s+(\w+)\s*%\}', texto_completo)
        
        # Buscar cualquier texto entre llaves
        cualquier_llave = re.findall(r'\{([^}]+)\}', texto_completo)
        
        print("=" * 60)
        print("PLACEHOLDERS ENCONTRADOS EN EL TEMPLATE")
        print("=" * 60)
        
        if placeholders_doble:
            print("\n[PLACEHOLDERS] Placeholders con {{variable}}:")
            for i, placeholder in enumerate(set(placeholders_doble), 1):
                print(f"   {i}. {{{{{placeholder}}}}}")
        else:
            print("\n[ADVERTENCIA] No se encontraron placeholders con {{ }}")
        
        if placeholders_dolar:
            print("\n[PLACEHOLDERS] Placeholders con {{$variable}}:")
            for i, placeholder in enumerate(set(placeholders_dolar), 1):
                print(f"   {i}. {{${placeholder}}}")
        else:
            print("\n[ADVERTENCIA] No se encontraron placeholders con {{$ }}")
        
        if placeholders_loop:
            print("\n[LOOPS] Loops encontrados ({% for %}):")
            for i, (var, lista) in enumerate(placeholders_loop, 1):
                print(f"   {i}. {{% for {var} in {lista} %}}")
        else:
            print("\n[ADVERTENCIA] No se encontraron loops con {% for %}")
        
        if cualquier_llave and not placeholders_doble and not placeholders_dolar:
            print("\n[INFO] Se encontraron textos entre llaves (pueden ser placeholders):")
            for i, texto in enumerate(set(cualquier_llave[:10]), 1):  # Mostrar solo los primeros 10
                print(f"   {i}. {{{{{texto}}}}}")
            if len(cualquier_llave) > 10:
                print(f"   ... y {len(cualquier_llave) - 10} mas")
        
        print("\n" + "=" * 60)
        print(f"Total placeholders {{}}: {len(set(placeholders_doble))}")
        print(f"Total placeholders {{$}}: {len(set(placeholders_dolar))}")
        print(f"Total loops: {len(placeholders_loop)}")
        print("=" * 60)
        
        if len(texto_completo) > 0:
            print(f"\n[INFO] Texto extraido del template ({len(texto_completo)} caracteres)")
            print("[INFO] Primeros 500 caracteres:")
            print("-" * 60)
            print(texto_completo[:500])
            print("-" * 60)
        
    except Exception as e:
        print(f"[ERROR] Error al leer el template: {e}")
        print("\n[INFO] Asegurate de que:")
        print("   1. El archivo existe en la ruta correcta")
        print("   2. El archivo no esta abierto en Word")
        print("   3. Tienes permisos de lectura")

if __name__ == "__main__":
    verificar_placeholders()

